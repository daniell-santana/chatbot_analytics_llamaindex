# embeddings.py - Vetorização do conhecimento
import os
import docx
from llama_index.core import Document, VectorStoreIndex, StorageContext, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb
from aux_objects import (
    CATEGORIAS_PRODUTOS, 
    KPIS_PADRAO,
    FORMAS_PAGAMENTO,
    STATUS_VENDAS,
    TEMPORADAS,
    FAIXAS_VALOR
)

# Configuração
Settings.embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Caminhos
CAMINHO_DOC_WORD = "data/sales_knowledge.docx"  # Crie este arquivo se quiser
DB_PATH = "./chromadb_sales"
COLLECTION_NAME = "sales_knowledge"

def criar_conhecimento_vendas():
    """Cria documentos de conhecimento para vendas."""
    documentos = []
    
    # 1. Categorias de produtos
    for categoria, info in CATEGORIAS_PRODUTOS.items():
        texto = f"""
        CATEGORIA: {categoria}
        Descrição: {info['descricao']}
        Subcategorias: {', '.join(info['subcategorias'])}
        
        Análises comuns para esta categoria:
        • Vendas sazonais: {categoria} tem picos em diferentes épocas
        • Margem típica: Varia entre 20-50% dependendo do produto
        • Ticket médio: Produtos de {categoria} costumam ter ticket entre R$ X e R$ Y
        """
        doc = Document(text=texto, metadata={"tipo": "categoria", "nome": categoria})
        documentos.append(doc)
    
    # 2. KPIs e Métricas
    kpis_texto = """
    KPIs PRINCIPAIS PARA ANÁLISE DE VENDAS:
    
    1. VENDAS BRUTAS: Soma de todos os valores de venda (valor_total)
       - Fórmula: df['valor_total'].sum()
       - Filtro: df['status'] == 'CONCLUÍDA'
    
    2. VENDAS LÍQUIDAS: Vendas brutas menos descontos e cancelamentos
       - Fórmula: df[df['status']=='CONCLUÍDA']['valor_total'].sum()
    
    3. TICKET MÉDIO: Valor médio por venda
       - Fórmula: df['valor_total'].mean()
       - Análise: Compare por região/categoria/forma pagamento
    
    4. MARGEM DE LUCRO: Porcentagem de lucro sobre vendas
       - Fórmula: (df['lucro'].sum() / df['valor_total'].sum()) * 100
       - Meta saudável: Acima de 20%
    
    5. CRESCIMENTO: Variação percentual entre períodos
       - Fórmula: ((Vendas atuais / Vendas período anterior) - 1) * 100
       - Análise: Mensal, trimestral, anual
    
    6. CUSTO DA MERCADORIA VENDIDA (CMV): Soma dos custos
       - Fórmula: df['custo_total'].sum()
    
    7. TAXA DE CANCELAMENTO: Vendas canceladas / Total vendas
       - Fórmula: (df[df['status']=='CANCELADA'].shape[0] / df.shape[0]) * 100
       - Alerta: Acima de 5% requer atenção
    """
    doc_kpis = Document(text=kpis_texto, metadata={"tipo": "kpis", "subtipo": "metricas"})
    documentos.append(doc_kpis)
    
    # 3. Formas de Pagamento
    pagamentos_texto = f"""
    FORMAS DE PAGAMENTO DISPONÍVEIS:
    {', '.join(FORMAS_PAGAMENTO)}
    
    Características:
    • Cartão Crédito: Maior ticket médio, possível parcelamento
    • Cartão Débito: Transação imediata, menor custo
    • PIX: Instantâneo, sem custo, em crescimento
    • Boleto: Prazo de pagamento, maior taxa de inadimplência
    • Dinheiro: Menos comum, usado em vendas presenciais
    
    Análise por forma de pagamento:
    - Ticket médio: Cartão Crédito > PIX > Débito
    - Taxa de conversão: PIX e Débito têm maior conversão
    - Custo operacional: Boleto > Cartão Crédito > PIX/Débito
    """
    doc_pagamentos = Document(text=pagamentos_texto, metadata={"tipo": "pagamentos"})
    documentos.append(doc_pagamentos)
    
    # 4. Análise Sazonal
    temporadas_texto = f"""
    SAZONALIDADE DAS VENDAS:
    
    Temporadas definidas:
    • ALTA TEMPORADA: {TEMPORADAS['ALTA']} (Natal, Ano Novo, Black Friday)
    • MÉDIA TEMPORADA: {TEMPORADAS['MEDIA']} (Férias de Julho)
    • BAIXA TEMPORADA: {TEMPORADAS['BAIXA']} (Pós-carnaval, Setembro)
    • NORMAL: {TEMPORADAS['NORMAL']} (Restante do ano)
    
    Comportamento por temporada:
    - Alta: Vendas aumentam 40-60%, descontos promocionais
    - Média: Vendas aumentam 20-30%, produtos sazonais
    - Baixa: Vendas caem 10-20%, foco em liquidação
    - Normal: Vendas estáveis, rotina operacional
    """
    doc_temporadas = Document(text=temporadas_texto, metadata={"tipo": "sazonalidade"})
    documentos.append(doc_temporadas)
    
    # 5. Regras de Negócio
    regras_texto = """
    REGRAS DE NEGÓCIO IMPORTANTES:
    
    1. Status de Vendas:
       • CONCLUÍDA: Venda finalizada e paga
       • CANCELADA: Venda cancelada antes da entrega
       • PENDENTE: Aguardando pagamento/confirmação
    
    2. Descontos:
       • Desconto padrão: Até 15%
       • Promoção especial: 16-30%
       • Liquidação: Acima de 30%
       • Descontos acima de 40% requerem aprovação
    
    3. Métricas de Performance:
       • Meta diária: Calculada com base no histórico
       • Meta mensal: Soma das metas diárias
       • Bônus por desempenho: Acima de 110% da meta
    
    4. Análise por Região:
       • Sudeste: Maior volume, concorrência alta
       • Sul: Ticket médio alto, fidelidade
       • Nordeste: Crescimento acelerado, sazonalidade forte
       • Centro-Oeste/Norte: Mercado em expansão
    """
    doc_regras = Document(text=regras_texto, metadata={"tipo": "regras_negocio"})
    documentos.append(doc_regras)
    
    # 6. Dicas de Análise
    dicas_texto = """
    DICAS PARA ANÁLISES EFETIVAS:
    
    1. Sempre filtre por status 'CONCLUÍDA' para métricas financeiras
    2. Compare períodos similares (ex: janeiro 2023 vs janeiro 2024)
    3. Segmentar por múltiplas dimensões (região + categoria + período)
    4. Observar outliers - valores muito altos/baixos merecem investigação
    5. Contextualizar números absolutos com porcentagens
    6. Considerar eventos externos (feriados, promoções, economia)
    
    PERGUNTAS FREQUENTES:
    • "Quais produtos têm maior margem?" → Filtre por margem_lucro
    • "Qual região cresce mais?" → Calcule crescimento por região
    • "Qual melhor forma de pagamento?" → Analise ticket médio e conversão
    • "Quando temos mais vendas?" → Agrupe por mês/trimestre
    • "Quem são os melhores vendedores?" → Agrupe por vendedor + valor_total
    """
    doc_dicas = Document(text=dicas_texto, metadata={"tipo": "dicas_analise"})
    documentos.append(doc_dicas)
    
    return documentos

def main():
    """Processo principal de ingestão."""
    print("🧠 Iniciando ingestão de conhecimento de vendas...")
    
    try:
        # 1. Cria documentos
        documentos = criar_conhecimento_vendas()
        
        # 2. Inicializa ChromaDB
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        
        # 3. Limpa coleção existente (se necessário)
        try:
            chroma_client.delete_collection(name=COLLECTION_NAME)
            print(f"🧹 Coleção '{COLLECTION_NAME}' limpa")
        except:
            pass  # Coleção não existia
        
        # 4. Cria nova coleção
        chroma_collection = chroma_client.get_or_create_collection(COLLECTION_NAME)
        vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        
        # 5. Cria índice
        index = VectorStoreIndex.from_documents(
            documentos, 
            storage_context=storage_context,
        )
        
        print(f"✅ Ingestão concluída com sucesso!")
        print(f"📚 Total de {len(documentos)} documentos vetorizados")
        print(f"💾 Salvos em: '{DB_PATH}'")
        print(f"🔍 Coleção: '{COLLECTION_NAME}'")
        
        # Resumo dos documentos
        print("\n📋 Documentos incluídos:")
        tipos = {}
        for doc in documentos:
            tipo = doc.metadata.get('tipo', 'desconhecido')
            tipos[tipo] = tipos.get(tipo, 0) + 1
        
        for tipo, quantidade in tipos.items():
            print(f"  • {tipo}: {quantidade} documento(s)")
        
    except Exception as e:
        print(f"❌ ERRO durante a ingestão: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()